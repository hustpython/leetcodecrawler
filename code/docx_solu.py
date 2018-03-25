# -*- coding:utf-8 -*-
'''
Created on  
2017年6月27日   下午9:52:43
@author: yzw
'''
import json
import docx
import sys
reload(sys) 
sys.setdefaultencoding('utf-8')
doc = docx.Document()
f=open('leetcode_solu1111.txt','r')              
for line in f.readlines():                                 
    data =unicode(line,'utf-8')
    
    try:
        if data.startswith('***'):
            #print'uu'
            con =data
            p=doc.add_paragraph()
            p.add_run(con).italic = True
        if data.startswith('Python_solution') or data.startswith('Best_solution'):
            con = data
            p=doc.add_paragraph()
            p.add_run(con).bold = True
            
        elif len(data)!=0 and not data.startswith('***'):    
            #print(data)     
            con = data
            doc.add_paragraph(con)
    except ValueError as e:
        print(data)

    
f.close()   
doc.save('solu3.docx')  