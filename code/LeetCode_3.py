# -*- coding:utf-8 -*-
'''
Created on  
2017年6月27日   下午9:52:43
@author: yzw
'''
import json
import docx
import sys
reload(sys) 
sys.setdefaultencoding('utf-8')
doc = docx.Document()
f=open('leetcode_database.txt','r')              
for line in f.readlines():                                 
    data=line.strip()
    data =unicode(data,'utf-8')
    try:
        if data == "'":
            data =''
        if data.endswith('***:'):
            con =data
            doc.add_heading(con, level=1) 
            doc.add_paragraph()
        if data.startswith('Example') or data.startswith('Note'):
            con = data
            p=doc.add_paragraph()
            p.add_run(con).bold = True
            
        elif len(data)!=0 and not data.endswith('***:'):         
            con = data
            doc.add_paragraph(con)
    except ValueError as e:
        print(data)

    
f.close()   
doc.save('exe.docx')  